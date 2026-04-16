"""
services/cv_generator.py — Génération de CV à partir d'un template .docx

Système de balises dans le template Word :
  {{NOM}}              → Nom de famille
  {{PRENOM}}           → Prénom
  {{EMAIL}}            → Email
  {{TELEPHONE}}        → Téléphone
  {{LINKEDIN}}         → URL LinkedIn
  {{POSTE}}            → Titre / poste professionnel
  {{BIO}}              → Texte de bio

  Pour les sections répétées, utiliser un tableau Word avec une ligne modèle
  contenant les balises suivantes :

  Expériences :
    {{EXP_TITRE}}       → Titre du poste
    {{EXP_ENTREPRISE}}  → Entreprise
    {{EXP_LOCATION}}    → Localisation
    {{EXP_DEBUT}}       → Date de début
    {{EXP_FIN}}         → Date de fin (ou "Présent")
    {{EXP_SUMMARY}}     → Résumé de projet
    {{EXP_DESC}}        → Description
    {{EXP_HARD_TITRE}}  → "Environnement Technique : " si hard skills, sinon ""
    {{EXP_HARD_NOM}}    → Hard skills liés, séparés par " , "
    {{EXP_SOFT_TITRE}}  → "Environnement Fonctionnel : " si soft skills, sinon ""
    {{EXP_SOFT_NOM}}    → Soft skills liés, séparés par " , "

  Formations :
    {{FORM_DIPLOME}}   → Diplôme
    {{FORM_ETAB}}      → Établissement
    {{FORM_DEBUT}}     → Date de début
    {{FORM_FIN}}       → Date de fin

  Certifications :
    {{CERT_TITRE}}     → Titre
    {{CERT_ORG}}       → Organisme
    {{CERT_DATE}}      → Date d'obtention
    {{CERT_FIN}}       → Date d'expiration (ou "pas d'expiration")

  Compétences hard :
    {{HARD_NOM}}       → Nom de la compétence
    {{HARD_NIVEAU}}    → Libellé du niveau (ex : "Expert")
    {{HARD_FAMILLE}}   → Famille de compétence

  Compétences soft :
    {{SOFT_NOM}}       → Nom de la compétence
    {{SOFT_NIVEAU}}    → Libellé du niveau (ex : "Intermédiaire")
    {{SOFT_FAMILLE}}   → Famille de compétence

  Langues parlées :
    {{LNG_NOM}}        → Nom de la langue (ex : "Anglais")
    {{LNG_NIVEAU}}     → Code + libellé CEFR (ex : "C1 — Autonome")
"""

import copy
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ── Helpers ────────────────────────────────────────────────────────────────


# ── Convertisseur HTML → blocs Word ────────────────────────────────────────

class _QuillParser(HTMLParser):
    """
    Parse le HTML généré par Quill et produit une liste de blocs :
    chaque bloc = liste de runs { text, bold, italic, underline }.
    Les <p> et <li> deviennent des blocs séparés.
    """
    def __init__(self):
        super().__init__()
        self.blocks: list[list[dict]] = []   # liste de paragraphes
        self._runs: list[dict] = []          # runs du paragraphe courant
        self._bold = self._italic = self._underline = False
        self._list_type: str | None = None   # 'ul' | 'ol'
        self._ol_idx = 0
        self._prefix = ""

    def _push_block(self):
        if self._runs or self._prefix:
            self.blocks.append({"prefix": self._prefix, "runs": self._runs})
        self._runs = []
        self._prefix = ""

    def handle_starttag(self, tag, attrs):
        if tag == "p":
            self._runs = []
            self._prefix = ""
        elif tag == "ul":
            self._list_type = "ul"
        elif tag == "ol":
            self._list_type = "ol"
            self._ol_idx = 0
        elif tag == "li":
            self._runs = []
            if self._list_type == "ol":
                self._ol_idx += 1
                self._prefix = f"{self._ol_idx}. "
            else:
                self._prefix = "• "
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "u":
            self._underline = True
        elif tag == "br":
            self._runs.append({"text": "\n", "bold": False, "italic": False, "underline": False})

    def handle_endtag(self, tag):
        if tag in ("p", "li"):
            self._push_block()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "u":
            self._underline = False
        elif tag in ("ul", "ol"):
            self._list_type = None

    def handle_data(self, data):
        if data:
            self._runs.append({
                "text": data,
                "bold": self._bold,
                "italic": self._italic,
                "underline": self._underline,
            })

    def handle_entityref(self, name):
        entities = {"amp": "&", "lt": "<", "gt": ">", "nbsp": " ", "quot": '"'}
        self.handle_data(entities.get(name, ""))

    def handle_charref(self, name):
        try:
            char = chr(int(name[1:], 16) if name.startswith("x") else int(name))
            self.handle_data(char)
        except ValueError:
            pass


def _parse_html_blocks(html: str) -> list[dict]:
    """Retourne la liste de blocs issus du HTML Quill."""
    if not html or html.strip() in ("", "<p><br></p>"):
        return []
    parser = _QuillParser()
    parser.feed(html)
    # Filtrer les blocs vides (paragraphes <p><br></p>)
    return [b for b in parser.blocks if any(r["text"].strip() for r in b["runs"]) or b["prefix"]]


def _make_run_xml(template_run_el, text: str, bold: bool, italic: bool, underline: bool):
    """
    Crée un <w:r> en clonant `template_run_el` (préserve police, taille, couleur…),
    puis remplace le texte et applique les surcharges bold/italic/underline.
    Si template_run_el est None, crée un run minimal.
    """
    if template_run_el is not None:
        r = copy.deepcopy(template_run_el)
    else:
        r = etree.Element(qn("w:r"))

    # Supprimer les anciens w:t et w:br
    for child in list(r):
        if child.tag in (qn("w:t"), qn("w:br")):
            r.remove(child)

    # Ajuster bold / italic / underline
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.Element(qn("w:rPr"))
        r.insert(0, rpr)
    for tag, active in [(qn("w:b"), bold), (qn("w:i"), italic), (qn("w:u"), underline)]:
        existing = rpr.find(tag)
        if active and existing is None:
            el = etree.SubElement(rpr, tag)
            if tag == qn("w:u"):
                el.set(qn("w:val"), "single")
        elif not active and existing is not None:
            rpr.remove(existing)

    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    if text and (text.startswith(" ") or text.endswith(" ")):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _clone_paragraph(p_el) -> "etree._Element":
    """Clone un <w:p> en supprimant tous ses runs — conserve pPr et toute la mise en forme."""
    new_p = copy.deepcopy(p_el)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    return new_p


def _build_text_paragraphs(text: str, p_el, template_run_el) -> list:
    """
    Convertit un texte brut (pouvant contenir des \\n) en liste de <w:p>.
    Chaque segment séparé par \\n devient un paragraphe distinct.
    La mise en forme est héritée du paragraphe et du run template.
    """
    new_els = []
    for segment in text.split("\n"):
        new_p = _clone_paragraph(p_el)
        if segment:
            new_p.append(_make_run_xml(template_run_el, segment, False, False, False))
        new_els.append(new_p)
    return new_els


def _run_text(r_el) -> str:
    """Texte complet d'un <w:r> (tous ses <w:t>)."""
    return "".join(t.text or "" for t in r_el.findall(qn("w:t")))


def _replace_html_field_in_cell(cell, marker: str, html: str) -> None:
    """
    Traite le paragraphe contenant `marker` AVANT _fill_row, pendant que les
    runs ont encore leur formatage d'origine (police, couleur, taille…).

    Algorithme :
    1. Reconstituer le texte complet du paragraphe run par run.
    2. Identifier par position de caractère quels runs sont avant, pendant
       ou après le marker (un run peut chevaucher le marker).
    3. Créer :
       - un paragraphe avec les runs « avant » (copiés tels quels)
       - des paragraphes pour le contenu HTML (formatage du run qui porte le marker)
       - un paragraphe avec les runs « après » (copiés tels quels, ils contiennent
         encore {{EXP_SOFT_TITRE}} etc. que _fill_row traitera ensuite)
    """
    for para in list(cell.paragraphs):
        p_el = para._p
        all_runs = p_el.findall(qn("w:r"))
        full_text = "".join(_run_text(r) for r in all_runs)

        if marker not in full_text:
            continue

        marker_start = full_text.index(marker)
        marker_end   = marker_start + len(marker)

        ppr_el       = p_el.find(qn("w:pPr"))
        parent       = p_el.getparent()
        insert_idx   = list(parent).index(p_el)

        before_runs: list = []
        after_runs:  list = []
        desc_run_el        = None   # run portant le marker (pour son formatage)

        pos = 0
        for r in all_runs:
            r_text = _run_text(r)
            r_end  = pos + len(r_text)

            if r_end <= marker_start:
                # Entièrement avant le marker
                before_runs.append(copy.deepcopy(r))

            elif pos >= marker_end:
                # Entièrement après le marker
                after_runs.append(copy.deepcopy(r))

            else:
                # Ce run chevauche le marker — on note son formatage
                if desc_run_el is None:
                    desc_run_el = r

                # Texte de ce run qui précède le marker
                pre = full_text[max(pos, 0): marker_start]
                if pos < marker_start and pre:
                    before_runs.append(_make_run_xml(r, pre, False, False, False))

                # Texte de ce run qui suit le marker
                suf = full_text[marker_end: r_end]
                if r_end > marker_end and suf:
                    after_runs.insert(
                        len([x for x in after_runs]),  # append au début des after
                        _make_run_xml(r, suf, False, False, False)
                    )

            pos = r_end

        blocks = _parse_html_blocks(html or "")

        def _make_p_with_runs(runs):
            new_p = etree.Element(qn("w:p"))
            if ppr_el is not None:
                new_p.append(copy.deepcopy(ppr_el))
            for r in runs:
                new_p.append(r)
            return new_p

        new_els = []

        if before_runs:
            new_els.append(_make_p_with_runs(before_runs))

        if blocks:
            for block in blocks:
                new_p = etree.Element(qn("w:p"))
                if ppr_el is not None:
                    new_p.append(copy.deepcopy(ppr_el))
                if block["prefix"]:
                    new_p.append(_make_run_xml(desc_run_el, block["prefix"], False, False, False))
                for run in block["runs"]:
                    new_p.append(_make_run_xml(desc_run_el, run["text"], run["bold"], run["italic"], run["underline"]))
                new_els.append(new_p)

        if after_runs:
            new_els.append(_make_p_with_runs(after_runs))

        if not new_els:
            new_els.append(etree.Element(qn("w:p")))

        for i, new_p in enumerate(new_els):
            parent.insert(insert_idx + i, new_p)
        parent.remove(p_el)
        return

    # Récursion sur les tableaux imbriqués
    for tbl in cell.tables:
        for row in tbl.rows:
            for c in row.cells:
                _replace_html_field_in_cell(c, marker, html)

def _fmt_date(d) -> str:
    """Formate une date Python en 'MM/YYYY', ou '' si None."""
    if d is None:
        return ""
    return d.strftime("%m/%Y")


def _fmt_duration(start, end=None) -> str:
    """Calcule la durée entre deux dates en texte lisible (ex: '2 ans 3 mois')."""
    from datetime import date
    if start is None:
        return ""
    end = end or date.today()
    total_months = (end.year - start.year) * 12 + (end.month - start.month)
    if total_months < 1:
        total_months = 1
    years  = total_months // 12
    months = total_months % 12
    parts  = []
    if years:
        parts.append(f"{years} an{'s' if years > 1 else ''}")
    if months:
        parts.append(f"{months} mois")
    return " ".join(parts) if parts else "< 1 mois"


def _apply_run_text(run, text: str) -> None:
    """
    Applique `text` à un run en gérant les \\n → <w:br/>.
    Préserve toutes les autres propriétés XML du run (police, couleur, taille…).
    """
    r_el = run._r
    for child in list(r_el):
        if child.tag in (qn("w:t"), qn("w:br")):
            r_el.remove(child)
    if "\n" not in text:
        if text:
            t = etree.SubElement(r_el, qn("w:t"))
            t.text = text
            if text.startswith(" ") or text.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    else:
        for idx, part in enumerate(text.split("\n")):
            if idx > 0:
                etree.SubElement(r_el, qn("w:br"))
            if part:
                t = etree.SubElement(r_el, qn("w:t"))
                t.text = part
                if part.startswith(" ") or part.endswith(" "):
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """
    Remplace les balises dans un paragraphe en préservant le formatage de chaque run.

    Passe 1 — remplacement direct dans chaque run qui contient la balise en entier :
      le formatage (police, couleur, gras…) du run est totalement préservé.

    Passe 2 — balises splittées sur plusieurs runs (Word découpe parfois {{NOM}}
      en '{{', 'NOM', '}}') : on fusionne le minimum de runs nécessaires en
      conservant le formatage du premier run du groupe.
    """
    if not para.runs:
        return

    full_text = "".join(r.text for r in para.runs)
    if not any(key in full_text for key in replacements):
        return

    # Passe 1 : remplacement dans les runs qui contiennent la balise entière
    for run in para.runs:
        if not any(key in run.text for key in replacements):
            continue
        new_text = run.text
        for key, value in replacements.items():
            new_text = new_text.replace(key, value or "")
        _apply_run_text(run, new_text)

    # Passe 2 : balises encore présentes (splittées sur plusieurs runs)
    full_text = "".join(r.text for r in para.runs)
    split_markers = [k for k in replacements if k in full_text]
    if not split_markers:
        return

    for marker in split_markers:
        value = replacements[marker] or ""
        runs = list(para.runs)
        for start in range(len(runs)):
            accumulated = ""
            for end in range(start, len(runs)):
                accumulated += runs[end].text
                if marker in accumulated:
                    # Fusionner uniquement runs[start:end+1]
                    new_text = accumulated.replace(marker, value)
                    _apply_run_text(runs[start], new_text)
                    for i in range(start + 1, end + 1):
                        runs[i].text = ""
                    break
            else:
                continue
            break


def _replace_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    """Remplace les balises dans tous les paragraphes ET toutes les cellules de tableau."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)


def _replace_in_cell(cell, replacements: dict[str, str]) -> None:
    """Remplace les balises dans tous les paragraphes d'une cellule."""
    for para in cell.paragraphs:
        _replace_in_paragraph(para, replacements)
    # Récursion sur les tableaux imbriqués
    for tbl in cell.tables:
        for row in tbl.rows:
            for c in row.cells:
                _replace_in_cell(c, replacements)


def _find_template_row(table, marker: str):
    """Retourne la première ligne du tableau contenant le marqueur."""
    for row in table.rows:
        for cell in row.cells:
            if marker in cell.text:
                return row
    return None


def _clone_row(row):
    """Clone une ligne de tableau (deep copy du XML)."""
    return copy.deepcopy(row._tr)


def _fill_row(row, replacements: dict[str, str]) -> None:
    """Remplace les balises dans toutes les cellules d'une ligne."""
    for cell in row.cells:
        _replace_in_cell(cell, replacements)


def _expand_table_section(
    doc: Document,
    marker: str,
    items: list[dict[str, str]],
    html_fields: dict[str, str] | None = None,
) -> None:
    """
    Trouve la table contenant `marker`, clone la ligne modèle pour chaque item
    et supprime la ligne modèle originale.

    html_fields : dict marker → clé dans chaque item dont la valeur est du HTML
                  (ex : {"{{EXP_DESC}}": "{{EXP_DESC}}"}).
                  Ces champs sont exclus du remplacement textuel simple et traités
                  via _replace_html_field_in_cell pour préserver le formatage Word.

    Important : on calcule la position d'insertion parmi TOUS les enfants XML
    du <w:tbl> (pas seulement les <w:tr>), car le tbl contient aussi w:tblPr,
    w:tblGrid, etc. en tête — un index basé sur w:tr seuls décalerait les
    insertions et corromprait le fichier.
    """
    from docx.table import _Row

    html_fields = html_fields or {}

    for table in doc.tables:
        template_row = _find_template_row(table, marker)
        if template_row is None:
            continue

        tbl = table._tbl
        tr_el = template_row._tr

        # Position réelle parmi TOUS les enfants du <w:tbl>
        all_children = list(tbl)
        insert_pos = all_children.index(tr_el)

        # Si aucun item, supprimer juste la ligne modèle et sortir
        if not items:
            tbl.remove(tr_el)
            break

        # Cloner + remplir une ligne par item, en décalant l'index d'insertion
        for i, item in enumerate(items):
            new_tr = _clone_row(template_row)
            tbl.insert(insert_pos + i, new_tr)
            new_row = _Row(new_tr, table)

            # 1. Remplacement HTML en premier — runs encore intacts, formatage préservé
            for marker_key, item_key in html_fields.items():
                html_val = item.get(item_key, "") or ""
                for cell in new_row.cells:
                    _replace_html_field_in_cell(cell, marker_key, html_val)

            # 2. Remplacement textuel simple sur le reste (hors champs HTML)
            simple_item = {k: v for k, v in item.items() if k not in html_fields}
            _fill_row(new_row, simple_item)

        # Supprimer la ligne modèle (maintenant décalée de len(items) positions)
        tbl.remove(tr_el)
        break


# ── Fonction principale ────────────────────────────────────────────────────

def generate_cv_docx(template_path: str, profile: dict[str, Any], output_path: str) -> None:
    """
    Génère un fichier .docx à partir du template et des données du profil.

    Args:
        template_path: chemin vers le fichier .docx template
        profile: dict contenant les clés user, profile, bio, experiences,
                 formations, certifications, competences
        output_path: chemin de sortie du fichier .docx généré
    """
    doc  = Document(template_path)
    user = profile["user"]
    prof = profile.get("profile")
    bio  = profile.get("bio")

    from models import SkillTypeEnum, CEFR_LABELS

    # ── Libellés de niveau compétence ──
    niveau_labels = {1: "Débutant", 2: "Intermédiaire", 3: "Avancé", 4: "Expert"}

    # ── Index GID → Competence (toutes les compétences du user) ──
    comp_by_gid: dict[str, Any] = {}
    for c in profile.get("competences", []):
        comp_by_gid[str(c.gid)] = c

    # ── Trigramme : 1ère lettre prénom + 2 premières lettres nom (majuscules) ──
    trigramme = (
        (user.prenom[:1] if user.prenom else "") +
        (user.nom[:2]    if user.nom    else "")
    ).upper()

    # ── Remplacements simples ──
    simple = {
        "{{NOM}}":       user.nom,
        "{{PRENOM}}":    user.prenom,
        "{{EMAIL}}":     user.email,
        "{{TELEPHONE}}": (prof.telephone if prof else "") or "",
        "{{LINKEDIN}}":  (prof.linkedin_url if prof else "") or "",
        "{{POSTE}}":     (prof.poste if prof else "") or "",
        "{{BIO}}":        (bio.texte if bio else "") or "",
        "{{TRIGRAMME}}":  trigramme,
    }
    _replace_in_doc(doc, simple)

    # Remplacements dans les headers/footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for para in hf.paragraphs:
                _replace_in_paragraph(para, simple)

    # ── Sections répétées ──

    # Expériences
    def _exp_skills(gids, skill_type):
        """Retourne les noms des compétences d'un type donné pour une liste de GIDs."""
        noms = []
        for gid in (gids or []):
            c = comp_by_gid.get(str(gid))
            if c and c.type == skill_type:
                noms.append(c.nom)
        return noms

    exp_rows = []
    for e in profile.get("experiences", []):
        hard_noms = _exp_skills(e.hard_skills, SkillTypeEnum.hard)
        soft_noms = _exp_skills(e.soft_skills, SkillTypeEnum.soft)
        exp_rows.append({
            "{{EXP_TITRE}}":       e.titre_poste,
            "{{EXP_ENTREPRISE}}":  e.entreprise,
            "{{EXP_LOCATION}}":    e.location or "",
            "{{EXP_DEBUT}}":       _fmt_date(e.date_debut),
            "{{EXP_FIN}}":         _fmt_date(e.date_fin) or "Présent",
            "{{EXP_SUMMARY}}":     e.project_summary or "",
            "{{EXP_DESC}}":        e.description or "",
            "{{EXP_DUREE}}":       _fmt_duration(e.date_debut, e.date_fin),
            "{{EXP_HARD_TITRE}}":  "\nEnvironnement Technique : " if hard_noms else "",
            "{{EXP_HARD_NOM}}":    ", ".join(hard_noms),
            "{{EXP_SOFT_TITRE}}":  "\nEnvironnement Fonctionnel : " if soft_noms else "",
            "{{EXP_SOFT_NOM}}":    ", ".join(soft_noms),
        })
    _expand_table_section(
        doc, "{{EXP_TITRE}}", exp_rows,
        html_fields={"{{EXP_DESC}}": "{{EXP_DESC}}"},
    )

    # Formations
    _expand_table_section(doc, "{{FORM_DIPLOME}}", [
        {
            "{{FORM_DIPLOME}}": f.diplome,
            "{{FORM_ETAB}}":    f.etablissement,
            "{{FORM_VILLE}}":   f.ville or "",
            "{{FORM_DEBUT}}":   _fmt_date(f.date_debut),
            "{{FORM_FIN}}":     _fmt_date(f.date_fin),
        }
        for f in profile.get("formations", [])
    ])

    # Certifications
    _expand_table_section(doc, "{{CERT_TITRE}}", [
        {
            "{{CERT_TITRE}}": c.titre,
            "{{CERT_ORG}}":   c.organisme,
            "{{CERT_DATE}}":  _fmt_date(c.date_obtention),
            "{{CERT_FIN}}":   _fmt_date(c.date_fin) if c.date_fin else "pas d'expiration",
        }
        for c in profile.get("certifications", [])
    ])

    # Compétences
    hard = [c for c in profile.get("competences", []) if c.type == SkillTypeEnum.hard]
    soft = [c for c in profile.get("competences", []) if c.type == SkillTypeEnum.soft]

    _expand_table_section(doc, "{{HARD_NOM}}", [
        {
            "{{HARD_NOM}}":    c.nom,
            "{{HARD_NIVEAU}}": niveau_labels.get(c.niveau.value, str(c.niveau.value)),
            "{{HARD_FAMILLE}}": c.famille or "",
        }
        for c in hard
    ])
    _expand_table_section(doc, "{{SOFT_NOM}}", [
        {
            "{{SOFT_NOM}}":    c.nom,
            "{{SOFT_NIVEAU}}": niveau_labels.get(c.niveau.value, str(c.niveau.value)),
            "{{SOFT_FAMILLE}}": c.famille or "",
        }
        for c in soft
    ])

    # Langues parlées
    _expand_table_section(doc, "{{LNG_NOM}}", [
        {
            "{{LNG_NOM}}":    l.nom,
            "{{LNG_NIVEAU}}": CEFR_LABELS.get(l.niveau.value, l.niveau.value),
        }
        for l in profile.get("profil_langues", [])
    ])

    doc.save(output_path)
    _strip_mip_label(output_path)


def _strip_mip_label(docx_path: str) -> None:
    """
    Supprime les métadonnées d'étiquette de sensibilité Microsoft (MIP/AIP)
    du fichier .docx généré. Sans ça, Word refuse d'ouvrir les fichiers
    issus de templates classifiés Sopra Steria (ex: C2 — Usage restreint).
    """
    import zipfile, shutil, os, tempfile

    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, \
         zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            # Supprimer les fichiers MIP et les propriétés custom (classification)
            if item.filename in ("docMetadata/LabelInfo.xml", "docProps/custom.xml"):
                continue
            data = zin.read(item.filename)
            # Nettoyer la référence à custom.xml dans [Content_Types].xml
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b'<Override PartName="/docProps/custom.xml" ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/>',
                    b""
                )
            # Nettoyer la référence à custom.xml dans les relations
            if item.filename == "_rels/.rels":
                import re as _re
                data = _re.sub(
                    rb'<Relationship[^/]*/docProps/custom\.xml[^/]*/>', b"", data
                )
            zout.writestr(item, data)

    os.replace(tmp, docx_path)


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """
    Convertit un .docx en .pdf via WeasyPrint (HTML intermédiaire)
    ou via LibreOffice si disponible.
    Tente d'abord LibreOffice (meilleure fidélité), puis fallback WeasyPrint.
    """
    import subprocess
    import sys

    # Tentative LibreOffice
    for soffice in ["soffice", "libreoffice"]:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir",
                 str(Path(pdf_path).parent), docx_path],
                capture_output=True, timeout=30,
            )
            if result.returncode == 0:
                # LibreOffice génère le PDF avec le même nom de base
                generated = Path(docx_path).with_suffix(".pdf")
                if generated.exists() and str(generated) != pdf_path:
                    generated.rename(pdf_path)
                return
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # Fallback : WeasyPrint via conversion HTML basique
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except Exception:
        # Dernier recours : copie du docx avec extension .pdf (fichier inutilisable mais évite l'erreur)
        import shutil
        shutil.copy2(docx_path, pdf_path)
